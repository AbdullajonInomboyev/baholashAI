import os
import uuid
import base64

import bleach
from docx import Document
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from lxml import etree

from qbank.models import (
    QuestionBank,
    QuestionGroup,
    Question,
    QuestionImage,
    TestBank,
    TestQuestion,
)
from assessment.services.formula_converter import OMMLToMathMLConverter


XML_NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v": "urn:schemas-microsoft-com:vml",
}


class DocxParseError(Exception):
    pass


class BaseDocxParser:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)
        self.related_parts = self.document.part.related_parts
        self.formula_converter = OMMLToMathMLConverter()

    def _clean_text(self, text):
        cleaned = (text or "").strip()
        return bleach.clean(
            cleaned,
            tags=["b", "i", "br", "strong", "em", "u", "sub", "sup"],
            strip=True,
        )

    def _xml_to_string(self, element):
        if element is None:
            return ""
        return etree.tostring(element, encoding="unicode")

    def _extract_image_by_rid(self, rid):
        part = self.related_parts.get(rid)
        if not part:
            return None

        filename = os.path.basename(str(part.partname))
        content_type = getattr(part, "content_type", "application/octet-stream")
        blob = part.blob

        return {
            "filename": filename,
            "content_type": content_type,
            "bytes": blob,
        }

    def _rid_to_data_url(self, rid):
        """rId -> brauzerbop PNG data URL. WMF/EMF bo'lsa konvertatsiya qiladi."""
        info = self._extract_image_by_rid(rid)
        if not info or not info.get("bytes"):
            return ""
        name = (info.get("filename") or "").lower()
        ct = (info.get("content_type") or "").lower()
        import base64 as _b64
        if name.endswith((".wmf", ".emf")) or "wmf" in ct or "emf" in ct or "x-emf" in ct:
            from qbank.services.formula_image import wmf_bytes_to_data_url
            ext = ".emf" if (name.endswith(".emf") or "emf" in ct) else ".wmf"
            return wmf_bytes_to_data_url(info["bytes"], ext=ext) or ""
        # oddiy raster (png/jpg/gif)
        b64 = _b64.b64encode(info["bytes"]).decode("ascii")
        ctype = ct if ct.startswith("image/") else "image/png"
        return f"data:{ctype};base64,{b64}"

    def _cell_image_data_url(self, cell):
        """Katak ichidagi BIRINCHI rasm/formulani (a:blip yoki v:imagedata=MathType)
        brauzerbop PNG data URL sifatida qaytaradi. Bo'lmasa ''."""
        tc = cell._tc
        R = XML_NAMESPACES["r"]
        # 1) DrawingML rasm (a:blip r:embed)
        for blip in tc.findall(f".//{{{XML_NAMESPACES['a']}}}blip"):
            rid = blip.get(f"{{{R}}}embed")
            if rid:
                url = self._rid_to_data_url(rid)
                if url:
                    return url
        # 2) VML rasm / MathType formula ko'rinishi (v:imagedata r:id)
        for imgd in tc.findall(f".//{{{XML_NAMESPACES['v']}}}imagedata"):
            rid = imgd.get(f"{{{R}}}id")
            if rid:
                url = self._rid_to_data_url(rid)
                if url:
                    return url
        return ""

    def _extract_paragraph_content(self, paragraph):
        """
        Paragraph ichidan:
        - text
        - formula (OMML)
        - image
        elementlarini bloklar ko‘rinishida qaytaradi.
        """
        blocks = []
        p_element = paragraph._element

        for child in p_element.iterchildren():
            tag = child.tag

            if tag.endswith("}r"):
                run_text_parts = []

                texts = child.findall(".//w:t", namespaces=XML_NAMESPACES)
                for t in texts:
                    if t.text:
                        run_text_parts.append(t.text)

                run_text = "".join(run_text_parts).strip()
                if run_text:
                    blocks.append({
                        "type": "text",
                        "value": run_text,
                    })

                omath_elements = child.findall(".//m:oMath", namespaces=XML_NAMESPACES)
                for omath in omath_elements:
                    formula_xml = self._xml_to_string(omath)
                    if formula_xml:
                        blocks.append({
                            "type": "formula_omml",
                            "value": formula_xml,
                        })

                omath_para_elements = child.findall(".//m:oMathPara", namespaces=XML_NAMESPACES)
                for omath_para in omath_para_elements:
                    formula_xml = self._xml_to_string(omath_para)
                    if formula_xml:
                        blocks.append({
                            "type": "formula_omml",
                            "value": formula_xml,
                        })

                blips = child.findall(".//a:blip", namespaces=XML_NAMESPACES)
                for blip in blips:
                    embed_rid = blip.get(f'{{{XML_NAMESPACES["r"]}}}embed')
                    if not embed_rid:
                        continue

                    image_info = self._extract_image_by_rid(embed_rid)
                    if image_info:
                        blocks.append({
                            "type": "image",
                            "filename": image_info["filename"],
                            "content_type": image_info["content_type"],
                            "bytes": image_info["bytes"],
                        })

            elif tag.endswith("}oMath"):
                formula_xml = self._xml_to_string(child)
                if formula_xml:
                    blocks.append({
                        "type": "formula_omml",
                        "value": formula_xml,
                    })

            elif tag.endswith("}oMathPara"):
                formula_xml = self._xml_to_string(child)
                if formula_xml:
                    blocks.append({
                        "type": "formula_omml",
                        "value": formula_xml,
                    })

        return blocks

    def _extract_cell_content(self, cell):
        """
        Cell ichidagi barcha paragraphlarni parse qiladi.
        Bitta savolda bir nechta formula va bir nechta rasm bo‘lishini hisobga oladi.
        """
        all_blocks = []
        html_parts = []
        raw_text_parts = []
        formulas_omml = []
        formulas_mathml = []
        images = []

        for paragraph in cell.paragraphs:
            paragraph_blocks = self._extract_paragraph_content(paragraph)

            if not paragraph_blocks:
                txt = self._clean_text(paragraph.text)
                if txt:
                    paragraph_blocks = [{"type": "text", "value": txt}]

            for block in paragraph_blocks:
                block_type = block["type"]

                if block_type == "text":
                    value = self._clean_text(block["value"])
                    if value:
                        raw_text_parts.append(value)
                        html_parts.append(f"<p>{value}</p>")
                        all_blocks.append({
                            "type": "text",
                            "value": value,
                        })

                elif block_type == "formula_omml":
                    value = block["value"]
                    if value:
                        try:
                            mathml = self.formula_converter.convert(value)
                        except Exception:
                            mathml = ""
                        formulas_omml.append(value)

                        if mathml:
                            formulas_mathml.append(mathml)
                            html_parts.append(mathml)
                        else:
                            html_parts.append("<p>[FORMULA]</p>")

                        all_blocks.append({
                            "type": "formula",
                            "omml": value,
                        })

                elif block_type == "image":
                    images.append(block)
                    html_parts.append("<p>[IMAGE]</p>")
                    all_blocks.append({
                        "type": "image",
                        "filename": block["filename"],
                        "content_type": block["content_type"],
                    })

        return {
            "content_html": "\n".join(html_parts),
            "raw_text": "\n".join(raw_text_parts),
            "formula_omml": "\n".join(formulas_omml),
            "formula_mathml": "\n".join(formulas_mathml),
            "raw_content_json": {
                "blocks": all_blocks,
            },
            "images": images,
        }

    def _save_question_images(self, question, images, prefix="question"):
        """
        Question uchun:
        - barcha rasmlarni QuestionImage modelga saqlaydi
        - birinchi rasmni question.image ga ham biriktiradi
        - JSON uchun URL/path qaytaradi
        """
        saved_urls = []

        for index, img in enumerate(images, start=1):
            original_name = img["filename"]
            ext = os.path.splitext(original_name)[1] or ".png"
            unique_name = f"{prefix}_{uuid.uuid4().hex}_{index}{ext}"

            django_file = ContentFile(img["bytes"], name=unique_name)

            question_image = QuestionImage.objects.create(
                question=question,
                image=django_file,
            )

            saved_urls.append(question_image.image.name)

            if not question.image:
                question.image.name = question_image.image.name

        return saved_urls

    def _save_test_question_images(self, test_question, images, upload_dir="test_questions", prefix="test_question"):
        """
        TestQuestion uchun:
        - hozircha alohida TestQuestionImage modeli yo‘q
        - barcha rasmlarni storage ga saqlaydi
        - birinchi rasmni test_question.image ga biriktiradi
        - JSON uchun URL/path qaytaradi
        """
        saved_urls = []

        for index, img in enumerate(images, start=1):
            original_name = img["filename"]
            ext = os.path.splitext(original_name)[1] or ".png"
            unique_name = f"{prefix}_{uuid.uuid4().hex}_{index}{ext}"
            relative_path = os.path.join(upload_dir, unique_name).replace("\\", "/")

            django_file = ContentFile(img["bytes"], name=unique_name)

            if index == 1 and not test_question.image:
                test_question.image.save(unique_name, django_file, save=False)
                saved_urls.append(test_question.image.name)
            else:
                saved_path = default_storage.save(relative_path, django_file)
                saved_urls.append(saved_path)

        return saved_urls

    def _first_image_data_url(self, images):
        """Birinchi rasmni base64 data URL ko‘rinishiga o‘tkazadi (bazada saqlanadi,
        Render'da deploydan keyin ham yo‘qolmaydi — S3 shart emas)."""
        if not images:
            return ""
        img = images[0]
        try:
            b64 = base64.b64encode(img["bytes"]).decode("ascii")
        except Exception:
            return ""
        ct = img.get("content_type") or "image/png"
        return f"data:{ct};base64,{b64}"

    def _inject_image_urls_into_blocks(self, raw_content_json, image_urls):
        """
        raw_content_json ichidagi image blocklarga url qo‘shadi.
        """
        blocks = raw_content_json.get("blocks", [])
        image_index = 0

        for block in blocks:
            if block.get("type") == "image" and image_index < len(image_urls):
                block["url"] = f"/media/{image_urls[image_index]}"
                image_index += 1

        return raw_content_json


class WrittenQuestionDocxParser(BaseDocxParser):
    """
    Yozma ish savollar banki parseri.

    Qoidalar:
    - Faylda bir nechta jadval bo‘ladi
    - Har jadval kamida 2 ustunli bo‘ladi
    - 1-satr, 2-ustun => group name
    - 2-satrdan boshlab, 2-ustun => savol
    """

    def parse_into_bank(self, bank: QuestionBank):
        if not isinstance(bank, QuestionBank):
            raise DocxParseError("bank QuestionBank obyekt bo‘lishi kerak.")

        tables = self.document.tables
        if not tables:
            raise DocxParseError("Docx ichida jadval topilmadi.")

        created_groups = []

        for table_index, table in enumerate(tables, start=1):
            if len(table.rows) < 2:
                raise DocxParseError(
                    f"{table_index}-jadvalda kamida 2 satr bo‘lishi kerak."
                )

            first_row = table.rows[0]
            if len(first_row.cells) < 2:
                raise DocxParseError(
                    f"{table_index}-jadvalda kamida 2 ustun bo‘lishi kerak."
                )

            group_name = self._clean_text(first_row.cells[1].text)
            if not group_name:
                raise DocxParseError(
                    f"{table_index}-jadvalning 1-satr 2-ustunida guruh nomi topilmadi."
                )

            group = QuestionGroup.objects.create(
                bank=bank,
                name=group_name,
                order=table_index,
                pick_count=1,
            )
            created_groups.append(group)

            for row_index, row in enumerate(table.rows[1:], start=2):
                if len(row.cells) < 2:
                    continue

                question_cell = row.cells[1]
                parsed = self._extract_cell_content(question_cell)

                has_any_content = any([
                    parsed["content_html"],
                    parsed["raw_text"],
                    parsed["formula_omml"],
                    parsed["images"],
                ])
                if not has_any_content:
                    continue

                question = Question.objects.create(
                    group=group,
                    content_html=parsed["content_html"],
                    raw_text=parsed["raw_text"],
                    formula_omml=parsed["formula_omml"],
                    formula_mathml=parsed["formula_mathml"],
                    image_data=(self._cell_image_data_url(question_cell)
                                or self._first_image_data_url(parsed["images"])),
                    raw_content_json=parsed["raw_content_json"],
                    difficulty="medium",
                    accessible=True,
                    allowed_answer_types=["text"],
                )

                image_urls = []
                if parsed["images"]:
                    image_urls = self._save_question_images(
                        question=question,
                        images=parsed["images"],
                        prefix="question",
                    )

                question.raw_content_json = self._inject_image_urls_into_blocks(
                    question.raw_content_json,
                    image_urls,
                )
                question.save()

        return created_groups


class TestQuestionDocxParser(BaseDocxParser):
    """
    Test savollari parseri.

    Qoidalar:
    - Faylda 1 ta asosiy jadval bo‘ladi
    - Jadvalda 6 ustun bo‘ladi:
      tr | savol | togri_javob | muqobil1 | muqobil2 | muqobil3
    """

    def parse_into_bank(self, bank: TestBank):
        if not isinstance(bank, TestBank):
            raise DocxParseError("bank TestBank obyekt bo‘lishi kerak.")

        tables = self.document.tables
        if not tables:
            raise DocxParseError("Docx ichida test jadvali topilmadi.")

        table = tables[0]

        if len(table.rows) < 2:
            raise DocxParseError("Test jadvalida kamida 2 satr bo‘lishi kerak.")

        created_questions = []

        for row_index, row in enumerate(table.rows[1:], start=2):
            if len(row.cells) < 6:
                raise DocxParseError(
                    f"{row_index}-satrda 6 ta ustun bo‘lishi kerak."
                )

            question_parsed = self._extract_cell_content(row.cells[1])

            correct_answer = self._clean_text(row.cells[2].text)
            option1 = self._clean_text(row.cells[3].text)
            option2 = self._clean_text(row.cells[4].text)
            option3 = self._clean_text(row.cells[5].text)

            # Savol va variant kataklaridagi formulalar (MathType/WMF) -> PNG data URL
            q_img = self._cell_image_data_url(row.cells[1])
            correct_img = self._cell_image_data_url(row.cells[2])
            option1_img = self._cell_image_data_url(row.cells[3])
            option2_img = self._cell_image_data_url(row.cells[4])
            option3_img = self._cell_image_data_url(row.cells[5])

            has_any_content = any([
                question_parsed["content_html"],
                question_parsed["raw_text"],
                question_parsed["formula_omml"],
                question_parsed["images"],
                q_img, correct_img, option1_img, option2_img, option3_img,
            ])
            if not has_any_content:
                continue

            test_question = TestQuestion.objects.create(
                bank=bank,
                question_html=question_parsed["content_html"],
                raw_text=question_parsed["raw_text"],
                formula_omml=question_parsed["formula_omml"],
                formula_mathml=question_parsed["formula_mathml"],
                image_data=(q_img or self._first_image_data_url(question_parsed["images"])),
                raw_content_json=question_parsed["raw_content_json"],
                correct_answer=correct_answer,
                option1=option1,
                option2=option2,
                option3=option3,
                correct_img=correct_img,
                option1_img=option1_img,
                option2_img=option2_img,
                option3_img=option3_img,
                accessible=True,
            )

            image_urls = []
            if question_parsed["images"]:
                image_urls = self._save_test_question_images(
                    test_question=test_question,
                    images=question_parsed["images"],
                    upload_dir="test_questions",
                    prefix="test_question",
                )

            test_question.raw_content_json = self._inject_image_urls_into_blocks(
                test_question.raw_content_json,
                image_urls,
            )
            test_question.save()
            created_questions.append(test_question)

        return created_questions